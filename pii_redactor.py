"""Small Flask service for detecting and replacing PII in common documents."""
from __future__ import annotations

import gc
import json
import logging
import os
import re
import shutil
import subprocess
import uuid
import warnings
import xml.etree.ElementTree as ET
from collections import Counter, defaultdict
from dataclasses import dataclass
from pathlib import Path
from typing import Any

os.environ.setdefault("BLIS_NUM_THREADS", "1")
os.environ.setdefault("OMP_NUM_THREADS", "1")
os.environ.setdefault("OPENBLAS_NUM_THREADS", "1")

import pandas as pd
import spacy
import fitz
from docx import Document
from faker import Faker
from flask import Flask, jsonify, request, send_file
from werkzeug.utils import secure_filename

BASE_DIR = Path(__file__).resolve().parent
INPUT_DIR, OUTPUT_DIR, DATASETS_DIR = (BASE_DIR / name for name in ("input", "output", "datasets"))
ALLOWED_EXTENSIONS = {".docx", ".doc", ".pdf", ".txt", ".csv", ".xlsx", ".xls", ".json", ".xml"}
PII_TYPES = ("PERSON", "EMAIL", "PHONE", "COMPANY", "ADDRESS", "SSN", "CREDIT_CARD", "DOB", "IP_ADDRESS", "LINKEDIN", "GITHUB", "WEBSITE")
PDF_BATCH_PAGE_LIMIT = 20
DOCX_BATCH_TARGET_LIMIT = 1_000
TABLE_BATCH_CELL_LIMIT = 10_000
TABLE_CHUNK_ROWS = 1_000
TABLE_CHUNK_CELLS = 10_000
CSV_STREAMING_BYTES = 8 * 1024 * 1024
STREAM_DETECTION_CACHE_LIMIT = 512

LOGGER = logging.getLogger(__name__)
if not LOGGER.handlers:
    handler = logging.StreamHandler()
    handler.setFormatter(logging.Formatter("%(asctime)s %(levelname)s %(message)s"))
    LOGGER.addHandler(handler)
LOGGER.setLevel(logging.INFO)
LOGGER.propagate = False


@dataclass(frozen=True)
class Entity:
    start: int
    end: int
    label: str
    text: str


@dataclass
class DocumentData:
    path: Path
    kind: str
    text: str
    content: Any = None


class FileExtractor:
    """Reads supported files and exposes their textual content."""

    def read(self, path: Path, include_text: bool = True) -> DocumentData:
        kind = path.suffix.lower()
        if kind not in ALLOWED_EXTENSIONS:
            raise ValueError(f"Unsupported format: {kind}")
        readers = {".txt": self._text, ".docx": self._docx, ".pdf": self._pdf,
                   ".csv": self._table, ".xlsx": self._table, ".xls": self._table,
                   ".json": self._json, ".xml": self._xml, ".doc": self._legacy_doc}
        return readers[kind](path, include_text=include_text)

    def _text(self, path: Path, include_text: bool = True) -> DocumentData:
        text = path.read_text(encoding="utf-8", errors="replace")
        return DocumentData(path, ".txt", text, text)

    def _docx(self, path: Path, include_text: bool = True) -> DocumentData:
        document = Document(path)
        if not include_text:
            # The writer traverses the document directly, so retaining a second
            # complete text representation only increases peak memory.
            return DocumentData(path, ".docx", "", document)
        parts = [paragraph.text for paragraph in document.paragraphs]
        parts += [cell.text for table in document.tables for row in table.rows for cell in row.cells]
        return DocumentData(path, ".docx", "\n".join(parts), document)

    def _pdf(self, path: Path, include_text: bool = True) -> DocumentData:
        # PDF writing extracts text from each page below. Avoid doing a second,
        # otherwise unused, full-document extraction before that work starts.
        return DocumentData(path, ".pdf", "")

    def _table(self, path: Path, include_text: bool = True) -> DocumentData:
        if not include_text and path.suffix.lower() == ".csv" and path.stat().st_size > CSV_STREAMING_BYTES:
            # Large CSV files are read in writer-owned chunks instead of being
            # represented by both a DataFrame and a concatenated text string.
            return DocumentData(path, ".csv", "")
        table = pd.read_csv(path, dtype=str, keep_default_na=False) if path.suffix == ".csv" else pd.read_excel(path, dtype=str, keep_default_na=False)
        text = "\n".join(table.astype(str).stack()) if include_text else ""
        return DocumentData(path, path.suffix.lower(), text, table)

    def _json(self, path: Path, include_text: bool = True) -> DocumentData:
        content = json.loads(path.read_text(encoding="utf-8"))
        return DocumentData(path, ".json", self._values(content) if include_text else "", content)

    def _xml(self, path: Path, include_text: bool = True) -> DocumentData:
        root = ET.parse(path).getroot()
        text = "\n".join(value for node in root.iter() for value in (node.text, *node.attrib.values()) if value) if include_text else ""
        return DocumentData(path, ".xml", text, root)

    def _legacy_doc(self, path: Path, include_text: bool = True) -> DocumentData:
        converted = self._office_convert(path, "docx", path.parent)
        data = self._docx(converted, include_text=include_text)
        return DocumentData(path, ".doc", data.text, data.content)

    @staticmethod
    def _office_convert(source: Path, format_name: str, target_dir: Path) -> Path:
        office = shutil.which("soffice") or shutil.which("libreoffice")
        if not office:
            raise ValueError("Legacy .doc needs LibreOffice installed, or convert it to .docx first.")
        subprocess.run([office, "--headless", "--convert-to", format_name, "--outdir", str(target_dir), str(source)], check=True, capture_output=True)
        converted = target_dir / f"{source.stem}.{format_name.split(':')[0]}"
        if not converted.exists():
            raise ValueError("LibreOffice could not convert the legacy .doc file.")
        return converted

    @staticmethod
    def _values(value: Any) -> str:
        if isinstance(value, dict):
            return "\n".join(FileExtractor._values(item) for item in value.values())
        if isinstance(value, list):
            return "\n".join(FileExtractor._values(item) for item in value)
        return str(value) if value is not None else ""


class PIIDetector:
    """Combines fast structured patterns with spaCy contextual recognition."""

    PATTERNS = {
        "EMAIL": r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b",
        "PHONE": r"(?<!\w)(?:\+?\d{1,3}[ .-]?)?(?:\(?\d{3}\)?[ .-]?)?\d{3}[ .-]\d{4}(?!\w)",
        "SSN": r"\b\d{3}-\d{2}-\d{4}\b",
        "CREDIT_CARD": r"\b(?:\d[ -]?){13,16}\b",
        "DOB": r"\b(?:0?[1-9]|1[0-2])[/-](?:0?[1-9]|[12]\d|3[01])[/-](?:19|20)\d{2}\b",
        "IP_ADDRESS": r"\b(?:25[0-5]|2[0-4]\d|1?\d?\d)(?:\.(?:25[0-5]|2[0-4]\d|1?\d?\d)){3}\b",
    }
    NER_LABELS = {"PERSON": "PERSON", "ORG": "COMPANY", "GPE": "ADDRESS", "LOC": "ADDRESS", "FAC": "ADDRESS"}
    NER_EXCLUSIONS = {"SSN", "DOB", "IP"}
    MODEL_NAMES = ("en_core_web_trf", "en_core_web_lg", "en_core_web_sm")
    ENTITY_PRIORITY = {"SSN": 3, "CREDIT_CARD": 3, "EMAIL": 3, "PHONE": 3, "DOB": 3, "IP_ADDRESS": 3, "LINKEDIN": 3, "GITHUB": 3, "WEBSITE": 3, "ADDRESS": 2, "PERSON": 1, "COMPANY": 1}
    COMPANY_SUFFIXES = ("inc", "ltd", "llc", "plc", "pvt", "corp", "corporation", "company", "technologies", "technology", "solutions", "systems", "services", "bank", "insurance", "holdings", "enterprises", "consulting", "group")
    PERSON_STOPWORDS = {"algorithm", "learning", "programming", "database", "structures", "concepts", "engineering", "embeddings", "search", "analytics", "developer", "framework", "model", "models"}
    PERSON_SUFFIXES = {"md", "phd", "jr", "sr", "ii", "iii", "iv"}
    PERSON_TITLES = {"mr", "mrs", "ms", "dr", "prof"}
    COMPANY_STOPWORDS = {"api", "ai", "ml", "sql", "pl/sql", "rag", "tensorflow", "python", "java", "react", "machine learning", "deep learning", "data structures", "prompt engineering", "semantic search", "langchain", "langgraph", "cnn", "nlp", "dsa", "keras", "phone", "email", "address", "iban", "ssn", "dob", "date"}
    FIELD_LABELS = {
        "PERSON": r"name|customer|employee|patient|contact",
        "COMPANY": r"company|organisation|organization|employer|business|firm",
    }
    PROFILE_PATTERNS = {
        "LINKEDIN": r"(?:https?://)?(?:www\.)?linkedin\.com/in/([A-Za-z0-9][A-Za-z0-9_-]{1,99})(?=[/?#\s]|$)",
        "GITHUB": r"(?:https?://)?(?:www\.)?github\.com/([A-Za-z0-9][A-Za-z0-9-]{0,38})(?=[/?#\s]|$)",
    }
    WEBSITE_PATTERN = r"(?<![@\w])(?:https?://|www\.)[A-Za-z0-9][A-Za-z0-9.-]*\.(?:com|org|net|io|dev|co|ai|me|app)(?:/[^\s<>()]*)?"
    EXTRA_PHONE_PATTERN = r"(?<!\w)(?:\+?91[-.\s]?)?[6-9]\d{4}[-.\s]\d{5}(?!\w)"
    ADDRESS_PATTERNS = (
        r"\b\d{1,5}\s+[A-Za-z][\w.'-]*(?:\s+[A-Za-z][\w.'-]*){0,4},\s*[A-Za-z][\w.'-]*(?:\s+[A-Za-z][\w.'-]*){0,3},\s*[A-Z]{2}\s+\d{5}(?:-\d{4})?\b",
        r"\b\d{1,6}\s+(?:[A-Za-z0-9.'-]+\s+){0,6}(?:street|st\.?|road|rd\.?|avenue|ave\.?|boulevard|blvd\.?|lane|ln\.?|drive|dr\.?|court|ct\.?|parkway|pkwy\.?|place|pl\.?|circle|cir\.?|highway|hwy\.?)\b(?:\s*(?:apt\.?|apartment|suite|unit|#)\s*[\w-]+)?",
        r"\b\d{1,5}\s+[A-Za-z][\w.'-]*(?:\s+\d+)?\s+(?:apt\.?|apartment|suite|unit)\s*[\w-]+(?:\s*\n\s*[A-Za-z][A-Za-z .'-]*){1,3}\s*,?\s*[A-Z]{2}\s*\n?\s*[A-Za-z .'-]+\s+\d{4,6}\b",
    )

    def __init__(self) -> None:
        self.nlp, self.model_name = self._load_model()
        self._direct_rules = tuple(
            (label, re.compile(expression)) for label, expression in self.PATTERNS.items()
        )
        self._address_rules = tuple(
            re.compile(expression, re.IGNORECASE | re.MULTILINE)
            for expression in self.ADDRESS_PATTERNS
        )
        self._profile_rules = tuple(
            (label, re.compile(expression, re.IGNORECASE))
            for label, expression in self.PROFILE_PATTERNS.items()
        )

    def detect(self, text: str) -> list[Entity]:
        return self._post_process(self._candidates(text, self.nlp(text)), text)

    def detect_many(self, texts: list[str]) -> list[list[Entity]]:
        documents = self.nlp.pipe(texts, batch_size=64)
        return [self._post_process(self._candidates(text, document), text)
                for text, document in zip(texts, documents)]

    def _candidates(self, text: str, document: Any) -> list[Entity]:
        sources = (
            self._regex_entities(text),
            self._profile_entities(text),
            self._field_entities(text),
            self._header_entities(text),
            self._ner_from_doc(document, text),
        )
        return [entity for source in sources for entity in source]

    def _load_model(self) -> tuple[Any, str]:
        model_names = ("en_core_web_sm",) if os.environ.get("RENDER", "").lower() == "true" else self.MODEL_NAMES
        for name in model_names:
            try:
                return spacy.load(name, exclude=["tagger", "parser", "lemmatizer", "attribute_ruler"]), name
            except OSError:
                continue
        warnings.warn("No spaCy NER model is installed; contextual PII detection is unavailable.", RuntimeWarning)
        return spacy.blank("en"), "blank"

    def _regex_entities(self, text: str) -> list[Entity]:
        detected: list[Entity] = []
        for label, rule in self._direct_rules:
            detected.extend(Entity(hit.start(), hit.end(), label, hit.group()) for hit in rule.finditer(text))
        for rule in self._address_rules:
            detected.extend(Entity(hit.start(), hit.end(), "ADDRESS", hit.group()) for hit in rule.finditer(text))
        indian_numbers = re.finditer(self.EXTRA_PHONE_PATTERN, text)
        detected.extend(Entity(hit.start(), hit.end(), "PHONE", hit.group()) for hit in indian_numbers)
        return detected

    def _profile_entities(self, text: str) -> list[Entity]:
        detected = [
            Entity(hit.start(1), hit.end(1), label, hit.group(1))
            for label, rule in self._profile_rules
            for hit in rule.finditer(text)
        ]
        for hit in re.finditer(self.WEBSITE_PATTERN, text, re.IGNORECASE):
            address = hit.group().casefold()
            if "linkedin.com" not in address and "github.com" not in address:
                detected.append(Entity(hit.start(), hit.end(), "WEBSITE", hit.group()))
        return detected

    def _field_entities(self, text: str) -> list[Entity]:
        entities: list[Entity] = []
        for label, fields in self.FIELD_LABELS.items():
            declaration = re.compile(rf"^\s*(?:{fields})\s*:\s*(?P<value>[^\n]+)", re.IGNORECASE | re.MULTILINE)
            for match in declaration.finditer(text):
                raw_value = match.group("value")
                value = raw_value.strip()
                start = match.start("value") + len(raw_value) - len(raw_value.lstrip())
                candidate = Entity(start, start + len(value), label, value)
                if self._valid_contextual_entity(candidate, text):
                    entities.append(candidate)
        return entities

    def _header_entities(self, text: str) -> list[Entity]:
        first_line = next((match for match in re.finditer(r"[^\n]+", text) if match.group().strip()), None)
        if not first_line:
            return []
        value = first_line.group().strip()
        start = first_line.start() + len(first_line.group()) - len(first_line.group().lstrip())
        return [Entity(start, start + len(value), "PERSON", value)] if self._valid_header_name(value) else []

    def _ner_entities(self, text: str) -> list[Entity]:
        return self._ner_from_doc(self.nlp(text), text)

    def _ner_from_doc(self, document: Any, text: str) -> list[Entity]:
        entities = []
        for entity in document.ents:
            if entity.label_ not in self.NER_LABELS or entity.text.upper() in self.NER_EXCLUSIONS:
                continue
            candidate = Entity(entity.start_char, entity.end_char, self.NER_LABELS[entity.label_], entity.text)
            candidate = self._expand_field_value(candidate, text)
            if self._valid_contextual_entity(candidate, text):
                entities.append(candidate)
        return entities

    def _expand_field_value(self, entity: Entity, text: str) -> Entity:
        if entity.label not in self.FIELD_LABELS:
            return entity
        line_start = text.rfind("\n", 0, entity.start) + 1
        prefix = text[line_start:entity.start]
        match = re.search(rf"\b(?:{self.FIELD_LABELS[entity.label]})\s*:\s*$", prefix, re.I)
        if not match:
            return entity
        line_end = text.find("\n", entity.end)
        line_end = len(text) if line_end == -1 else line_end
        value_start = line_start + match.end()
        value = text[value_start:line_end]
        left, right = len(value) - len(value.lstrip()), len(value) - len(value.rstrip())
        return Entity(value_start + left, line_end - right, entity.label, value.strip())

    def _valid_contextual_entity(self, entity: Entity, text: str) -> bool:
        if entity.label == "PERSON":
            return self._valid_person(entity.text, text, entity.start)
        if entity.label == "COMPANY":
            return self._valid_company(entity.text, text, entity.start)
        return self._has_address_cue(text, entity.start, entity.end)

    def _valid_person(self, value: str, text: str, start: int) -> bool:
        tokens = value.split()
        if tokens and tokens[0].casefold().rstrip(".") in self.PERSON_TITLES:
            tokens = tokens[1:]
        name_tokens = tokens[:-1] if tokens and tokens[-1].casefold().rstrip(".") in self.PERSON_SUFFIXES else tokens
        if "\n" in value or any(not self._name_token(token) for token in name_tokens):
            return False
        if any(token.casefold().strip(".") in self.PERSON_STOPWORDS for token in name_tokens):
            return False
        if len(name_tokens) in (2, 3, 4):
            return True
        context = text[max(0, start - 35):start].casefold()
        return len(name_tokens) == 1 and bool(re.search(r"\b(name|called|contact|customer|employee|patient)\b", context))

    @staticmethod
    def _name_token(token: str) -> bool:
        return bool(re.fullmatch(r"(?:[A-ZÀ-ÖØ-Þ][A-Za-zÀ-ÖØ-öø-ÿ'’-]+|[A-Z]\.)", token))

    def _valid_company(self, value: str, text: str, start: int) -> bool:
        if "\n" in value:
            return False
        suffix = r"\b(" + "|".join(self.COMPANY_SUFFIXES) + r")\.?$"
        context = text[max(0, start - 45):start].casefold()
        context_cue = r"\b(company|organisation|organization|employer|business|firm|vendor|client|work(?:s|ing)?\s+(?:at|for)|employed\s+by)\b"
        words = set(re.findall(r"[a-z]+(?:/[a-z]+)?", value.casefold()))
        capitalised = all(word[:1].isupper() for word in value.replace("-", " ").replace(",", " ").split() if word)
        return bool(re.search(suffix, value, re.I) or re.search(context_cue, context) or (capitalised and not words & self.COMPANY_STOPWORDS))

    def _valid_header_name(self, value: str) -> bool:
        if any(character.isdigit() for character in value) or re.search(r"[:,;@/]|\b(engineer|developer|manager|analyst|consultant)\b", value, re.I):
            return False
        tokens = value.replace(".", "").split()
        return 1 <= len(tokens) <= 4 and all(token.isupper() or self._name_token(token.title()) for token in tokens)

    @staticmethod
    def _has_address_cue(text: str, start: int, end: int) -> bool:
        prefix = text[max(0, start - 25):start]
        if re.search(r"address\s+of\s*$", prefix, re.I):
            return False
        context = text[max(0, start - 40):min(len(text), end + 40)]
        return bool(re.search(r"\b(address|located|resides|lives|street|road|avenue|apt|suite|unit|postal|postcode|zip)\b|\d{4,6}", context, re.I))

    def _post_process(self, entities: list[Entity], text: str) -> list[Entity]:
        return self._merge_person_tokens(self._merge(entities), text)

    def _merge(self, entities: list[Entity]) -> list[Entity]:
        by_location = {(item.start, item.end, item.label): item for item in entities}
        ranked = sorted(
            by_location.values(),
            key=lambda item: (item.start, -(item.end - item.start), -self.ENTITY_PRIORITY[item.label]),
        )
        selected: list[Entity] = []
        for item in ranked:
            if not selected or item.start >= selected[-1].end:
                selected.append(item)
        return selected

    def _merge_person_tokens(self, entities: list[Entity], text: str) -> list[Entity]:
        resolved: list[Entity] = []
        for current in entities:
            previous = resolved[-1] if resolved else None
            if previous is None or previous.label != current.label or current.label != "PERSON":
                resolved.append(current)
                continue
            between = text[previous.end:current.start]
            if not between.isspace():
                resolved.append(current)
                continue
            joined = Entity(previous.start, current.end, "PERSON", text[previous.start:current.end])
            if self._valid_person(joined.text, text, joined.start):
                resolved[-1] = joined
            else:
                resolved.append(current)
        return resolved


class PIIReplacer:
    """Creates stable, realistic replacements for detected values."""

    def __init__(self) -> None:
        self.fake = Faker()
        self.mapping: dict[tuple[str, str], str] = {}
        self._streaming_entities: dict[str, list[Entity]] = {}
        self.generators = {
            "PERSON": self.fake.name, "EMAIL": self.fake.email, "PHONE": self.fake.phone_number,
            "COMPANY": self.fake.company, "ADDRESS": self.fake.address, "SSN": self.fake.ssn,
            "CREDIT_CARD": self.fake.credit_card_number, "DOB": self._date_of_birth,
            "IP_ADDRESS": self.fake.ipv4_private, "LINKEDIN": self.fake.user_name,
            "GITHUB": self.fake.user_name, "WEBSITE": self.fake.url,
        }

    def replace(self, text: str, entities: list[Entity]) -> str:
        for entity in reversed(entities):
            replacement = self._replacement(entity)
            text = text[:entity.start] + replacement + text[entity.end:]
        return text

    def redact_value(self, value: Any, detector: PIIDetector) -> Any:
        if not isinstance(value, str):
            return value
        # Low-memory writers may encounter repeated cell or paragraph values.
        # Keep only a small, bounded cache so they avoid rerunning NER without
        # retaining the document-wide value/entity lists used by batch mode.
        if len(value) <= 4_096 and value in self._streaming_entities:
            return self.replace(value, self._streaming_entities[value])
        entities = detector.detect(value)
        if len(value) <= 4_096:
            if len(self._streaming_entities) >= STREAM_DETECTION_CACHE_LIMIT:
                self._streaming_entities.clear()
            self._streaming_entities[value] = entities
        return self.replace(value, entities)

    def redact_many(self, values: list[str], detector: PIIDetector) -> list[str]:
        unique_values = list(dict.fromkeys(values))
        replacements = {
            value: self.replace(value, entities)
            for value, entities in zip(unique_values, detector.detect_many(unique_values))
        }
        return [replacements[value] for value in values]

    def _replacement(self, entity: Entity) -> str:
        key = (entity.label, entity.text.casefold())
        if key not in self.mapping:
            self.mapping[key] = self.generators[entity.label]().replace("\n", ", ")
        return self.mapping[key]

    def _date_of_birth(self) -> str:
        return self.fake.date_of_birth(minimum_age=18, maximum_age=80).strftime("%m/%d/%Y")


class FileWriter:
    """Writes redacted content while retaining the source file format."""

    PDF_COMPANY_SUFFIXES = ("inc", "ltd", "llc", "limited", "corp", "corporation", "company")
    PDF_NON_NAME_WORDS = {"algorithm", "learning", "programming", "structures", "concepts", "developed", "contributed", "implemented", "designed", "trained", "selected"}

    def write(self, document: DocumentData, detector: PIIDetector, replacer: PIIReplacer) -> Path:
        output = OUTPUT_DIR / f"redacted_{document.path.stem}_{uuid.uuid4().hex[:8]}{document.kind}"
        writers = {".txt": self._text, ".docx": self._docx, ".pdf": self._pdf, ".csv": self._table,
                   ".xlsx": self._table, ".xls": self._table, ".json": self._json, ".xml": self._xml, ".doc": self._legacy_doc}
        writers[document.kind](document, output, detector, replacer)
        return output

    @staticmethod
    def _text(document: DocumentData, output: Path, detector: PIIDetector, replacer: PIIReplacer) -> None:
        output.write_text(replacer.replace(document.text, detector.detect(document.text)), encoding="utf-8")

    @staticmethod
    def _docx(document: DocumentData, output: Path, detector: PIIDetector, replacer: PIIReplacer) -> None:
        paragraph_count = len(document.content.paragraphs)
        cell_count = sum(1 for table in document.content.tables for row in table.rows for cell in row.cells)
        target_count = paragraph_count + cell_count
        if target_count <= DOCX_BATCH_TARGET_LIMIT:
            LOGGER.info("Processing DOCX in fast batch mode (%s targets)", target_count)
            targets = list(document.content.paragraphs)
            targets += [cell for table in document.content.tables for row in table.rows for cell in row.cells]
            for target, text in zip(targets, replacer.redact_many([target.text for target in targets], detector)):
                target.text = text
        else:
            LOGGER.info("Processing DOCX in low-memory mode (%s targets)", target_count)
            for index, paragraph in enumerate(document.content.paragraphs, start=1):
                paragraph.text = replacer.redact_value(paragraph.text, detector)
                if index % TABLE_CHUNK_ROWS == 0:
                    LOGGER.info("Processed %s/%s DOCX paragraphs", index, paragraph_count)
                    gc.collect()
            gc.collect()
            cell_index = 0
            for table_number, table in enumerate(document.content.tables, start=1):
                for row in table.rows:
                    for cell in row.cells:
                        cell.text = replacer.redact_value(cell.text, detector)
                        cell_index += 1
                        if cell_index % TABLE_CHUNK_ROWS == 0:
                            LOGGER.info("Processed %s/%s DOCX cells", cell_index, cell_count)
                            gc.collect()
                LOGGER.info("Completed DOCX table %s", table_number)
                gc.collect()
        LOGGER.info("Saving DOCX")
        document.content.save(output)

    @staticmethod
    def _pdf(document: DocumentData, output: Path, detector: PIIDetector, replacer: PIIReplacer) -> None:
        pdf = fitz.open(document.path)
        try:
            page_count = len(pdf)
            LOGGER.info("Loading PDF with %s pages", page_count)
            if page_count <= PDF_BATCH_PAGE_LIMIT:
                LOGGER.info("Processing PDF in fast batch mode")
                page_texts = [page.get_text("text") for page in pdf]
                page_entities = detector.detect_many(page_texts)
                for page_number, (page, entities) in enumerate(zip(pdf, page_entities), start=1):
                    LOGGER.info("Processing page %s/%s", page_number, page_count)
                    FileWriter._apply_pdf_redactions(page, entities, replacer)
            else:
                LOGGER.info("Processing PDF in low-memory mode")
                for page_number in range(page_count):
                    page = page_text = entities = None
                    try:
                        LOGGER.info("Processing page %s/%s", page_number + 1, page_count)
                        page = pdf.load_page(page_number)
                        page_text = page.get_text("text")
                        entities = detector.detect(page_text)
                        FileWriter._apply_pdf_redactions(page, entities, replacer)
                    finally:
                        # Do not retain text, entities, or PyMuPDF page references
                        # after the page has been written.
                        del entities, page_text, page
                        gc.collect()
            LOGGER.info("Saving PDF")
            pdf.save(output, garbage=4, deflate=True)
            LOGGER.info("Completed PDF redaction successfully")
        finally:
            pdf.close()

    @staticmethod
    def _apply_pdf_redactions(page: Any, entities: list[Entity], replacer: PIIReplacer) -> None:
        replacements = FileWriter._pdf_replacements(page, entities, replacer)
        try:
            for rect, _ in replacements:
                page.add_redact_annot(rect, fill=(1, 1, 1), cross_out=False)
            if replacements:
                page.apply_redactions()
                for rect, value in replacements:
                    FileWriter._insert_pdf_text(page, rect, value)
        finally:
            del replacements

    @staticmethod
    def _pdf_replacements(page: Any, entities: list[Entity], replacer: PIIReplacer) -> list[tuple[Any, str]]:
        replacements = []
        seen = set()
        for entity in entities:
            text = FileWriter._safe_pdf_entity(entity)
            if not text:
                continue
            key = (entity.label, text.casefold())
            if key in seen:
                continue
            seen.add(key)
            value = replacer.replace(text, [Entity(0, len(text), entity.label, text)])
            replacements += [(rect, value) for rect in page.search_for(text)]
        return replacements

    @staticmethod
    def _insert_pdf_text(page: Any, rect: Any, value: str) -> None:
        size = max(6, rect.height * 0.8)
        while size > 5 and page.insert_textbox(rect, value, fontsize=size, fontname="helv", color=(0, 0, 0)) < 0:
            size -= 0.5

    @staticmethod
    def _safe_pdf_entity(entity: Entity) -> str | None:
        if entity.label in PIIDetector.PATTERNS or entity.label in {"LINKEDIN", "GITHUB", "WEBSITE"}:
            return entity.text
        text = entity.text.strip()
        if entity.label == "PERSON":
            text = text.splitlines()[0].strip()
            words = text.lower().split()
            if len(words) in (2, 3) and re.fullmatch(r"[A-Z][a-z]+(?:[-'][A-Za-z]+)?(?:\s+[A-Z][a-z]+(?:[-'][A-Za-z]+)?){1,2}", text) and not set(words) & FileWriter.PDF_NON_NAME_WORDS:
                return text
        if entity.label == "COMPANY" and "\n" not in text and re.search(rf"\b({'|'.join(FileWriter.PDF_COMPANY_SUFFIXES)})\.?$", text, re.I):
            return text
        if entity.label == "ADDRESS" and re.search(r"\b\d{1,5}\s+.+\b(street|st|road|rd|avenue|ave|lane|ln|boulevard|blvd)\b", text, re.I):
            return text
        return None

    @staticmethod
    def _table(document: DocumentData, output: Path, detector: PIIDetector, replacer: PIIReplacer) -> None:
        if document.kind == ".csv" and document.content is None:
            FileWriter._stream_csv(document.path, output, detector, replacer)
            return
        table = document.content
        if table.size <= TABLE_BATCH_CELL_LIMIT:
            LOGGER.info("Processing dataframe in fast batch mode (%s cells)", table.size)
            values = table.astype(str).to_numpy().ravel().tolist()
            redacted = replacer.redact_many(values, detector)
            rows = [redacted[index:index + len(table.columns)]
                    for index in range(0, len(redacted), len(table.columns))]
            table = pd.DataFrame(rows, columns=table.columns, index=table.index)
        else:
            LOGGER.info("Processing dataframe in low-memory mode (%s cells)", table.size)
            FileWriter._redact_table_incrementally(table, detector, replacer)
        if document.kind == ".csv":
            LOGGER.info("Saving CSV")
            table.to_csv(output, index=False)
        elif document.kind == ".xls":
            LOGGER.info("Saving XLS")
            import xlwt
            workbook = xlwt.Workbook()
            sheet = workbook.add_sheet("Redacted")
            for column, value in enumerate(table.columns):
                sheet.write(0, column, value)
            for row, values in enumerate(table.itertuples(index=False, name=None), start=1):
                for column, value in enumerate(values): sheet.write(row, column, value)
            workbook.save(str(output))
        else:
            LOGGER.info("Saving XLSX")
            table.to_excel(output, index=False)

    @staticmethod
    def _stream_csv(source: Path, output: Path, detector: PIIDetector, replacer: PIIReplacer) -> None:
        for chunk_number, table in enumerate(pd.read_csv(source, dtype=str, keep_default_na=False, chunksize=TABLE_CHUNK_ROWS), start=1):
            LOGGER.info("Processing CSV dataframe chunk %s", chunk_number)
            FileWriter._redact_table_incrementally(table, detector, replacer)
            table.to_csv(output, index=False, mode="w" if chunk_number == 1 else "a", header=chunk_number == 1)
            del table
            gc.collect()

    @staticmethod
    def _redact_table_incrementally(table: Any, detector: PIIDetector, replacer: PIIReplacer) -> None:
        columns = len(table.columns)
        rows_per_chunk = max(1, TABLE_CHUNK_CELLS // max(1, columns))
        for start in range(0, len(table.index), rows_per_chunk):
            end = min(start + rows_per_chunk, len(table.index))
            values = [str(table.iat[row, column])
                      for row in range(start, end)
                      for column in range(columns)]
            redacted = replacer.redact_many(values, detector)
            redacted_index = 0
            for row in range(start, end):
                for column in range(columns):
                    table.iat[row, column] = redacted[redacted_index]
                    redacted_index += 1
            LOGGER.info("Processed dataframe rows %s-%s", start + 1, end)
            del values, redacted
            gc.collect()

    def _json(self, document: DocumentData, output: Path, detector: PIIDetector, replacer: PIIReplacer) -> None:
        LOGGER.info("Saving JSON")
        output.write_text(json.dumps(self._walk(document.content, detector, replacer), indent=2), encoding="utf-8")

    def _xml(self, document: DocumentData, output: Path, detector: PIIDetector, replacer: PIIReplacer) -> None:
        for index, node in enumerate(document.content.iter(), start=1):
            node.text = replacer.redact_value(node.text, detector)
            for key, value in node.attrib.items():
                node.attrib[key] = replacer.redact_value(value, detector)
            if index % TABLE_CHUNK_ROWS == 0:
                LOGGER.info("Processed %s XML nodes", index)
                gc.collect()
        LOGGER.info("Saving XML")
        ET.ElementTree(document.content).write(output, encoding="utf-8", xml_declaration=True)

    def _legacy_doc(self, document: DocumentData, output: Path, detector: PIIDetector, replacer: PIIReplacer) -> None:
        temporary_docx = output.with_suffix(".docx")
        self._docx(document, temporary_docx, detector, replacer)
        FileExtractor._office_convert(temporary_docx, "doc:MS Word 97", output.parent)
        temporary_docx.unlink(missing_ok=True)

    def _walk(self, value: Any, detector: PIIDetector, replacer: PIIReplacer) -> Any:
        if isinstance(value, dict):
            for key, item in value.items():
                value[key] = self._walk(item, detector, replacer)
            return value
        if isinstance(value, list):
            for index, item in enumerate(value):
                value[index] = self._walk(item, detector, replacer)
            return value
        return replacer.redact_value(value, detector)


class Evaluator:
    """Scores exact text-and-type matches against labelled source documents."""

    def __init__(self, extractor: FileExtractor, detector: PIIDetector) -> None:
        self.extractor, self.detector = extractor, detector

    def evaluate(self, datasets: Path = DATASETS_DIR) -> dict[str, Any]:
        totals, by_type = defaultdict(int), defaultdict(lambda: defaultdict(int))
        label_files = [datasets / "labels.json"] if (datasets / "labels.json").exists() else datasets.glob("*/labels.json")
        texts, expectations = [], []

        def evaluate_batch() -> None:
            if not texts:
                return
            for expected, entities in zip(expectations, self.detector.detect_many(texts)):
                found = Counter((entity.label, self._normalise(entity.text)) for entity in entities)
                self._count(expected, found, totals, by_type)
                totals["documents"] += 1
            texts.clear()
            expectations.clear()

        for labels_path in label_files:
            for item in json.loads(labels_path.read_text(encoding="utf-8")):
                text = self.extractor.read(labels_path.parent / "original" / item["file"]).text
                expected = Counter((entity["label"], self._normalise(entity["text"])) for entity in item["entities"])
                texts.append(text)
                expectations.append(expected)
                if len(texts) == 64:
                    evaluate_batch()
        evaluate_batch()
        report = self._metrics(totals)
        report["by_type"] = {label: self._metrics(by_type[label]) for label in PII_TYPES}
        return report

    @staticmethod
    def _count(expected: Counter, found: Counter, totals: dict, by_type: dict) -> None:
        for key in expected.keys() | found.keys():
            label, _ = key
            tp = min(expected[key], found[key])
            fp, fn = found[key] - tp, expected[key] - tp
            totals["tp"] += tp; by_type[label]["tp"] += tp
            totals["fp"] += fp; by_type[label]["fp"] += fp
            totals["fn"] += fn; by_type[label]["fn"] += fn
        totals["entities"] += sum(expected.values())

    @staticmethod
    def _normalise(value: str) -> str:
        return " ".join(value.casefold().split())

    @staticmethod
    def _metrics(values: dict) -> dict[str, Any]:
        tp, fp, fn = values["tp"], values["fp"], values["fn"]
        precision = tp / (tp + fp) if tp + fp else 0
        recall = tp / (tp + fn) if tp + fn else 0
        return {"documents": values.get("documents", 0), "entities": values.get("entities", tp + fn), "true_positives": tp,
                "false_positives": fp, "false_negatives": fn, "precision": round(precision, 4), "recall": round(recall, 4),
                "f1_score": round(2 * precision * recall / (precision + recall), 4) if precision + recall else 0,
                "accuracy": round(tp / (tp + fp + fn), 4) if tp + fp + fn else 0}


app = Flask(__name__)
for directory in (INPUT_DIR, OUTPUT_DIR):
    directory.mkdir(parents=True, exist_ok=True)
extractor, detector, writer = FileExtractor(), PIIDetector(), FileWriter()


@app.get("/")
def home():
    return send_file(BASE_DIR / "index.html")


@app.post("/redact")
def redact():
    upload = request.files.get("file")
    if not upload or not upload.filename:
        return jsonify(error="Choose a file to redact."), 400
    filename = secure_filename(upload.filename)
    if Path(filename).suffix.lower() not in ALLOWED_EXTENSIONS:
        return jsonify(error="This file format is not supported."), 400
    source = INPUT_DIR / f"{uuid.uuid4().hex}_{filename}"
    try:
        LOGGER.info("Upload received (%s)", Path(filename).suffix.lower())
        upload.save(source)
        LOGGER.info("Reading document")
        document = extractor.read(source, include_text=False)
        result = writer.write(document, detector, PIIReplacer())
        response = send_file(result, as_attachment=True, download_name=f"redacted_{filename}")

        def remove_temporary_files() -> None:
            source.unlink(missing_ok=True)
            result.unlink(missing_ok=True)

        response.call_on_close(remove_temporary_files)
        return response
    except Exception as error:
        source.unlink(missing_ok=True)
        return jsonify(error=str(error)), 400


@app.post("/evaluate")
def evaluate():
    return jsonify(Evaluator(extractor, detector).evaluate())


if __name__ == "__main__":
    app.run(host="0.0.0.0", port=int(os.environ.get("PORT", "5000")))
